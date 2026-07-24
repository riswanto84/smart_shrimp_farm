import csv
import io
import json
from pathlib import Path

from django.conf import settings

MAX_EXTRACTED_CHARS = getattr(settings, 'CHAT_AI_MAX_EXTRACTED_CHARS', 50000)
ALLOWED_EXTENSIONS = {'.txt', '.md', '.csv', '.json', '.pdf', '.docx', '.xlsx', '.xls', '.png', '.jpg', '.jpeg', '.webp'}


def validate_upload(uploaded):
    ext = Path(uploaded.name).suffix.lower()
    max_bytes = getattr(settings, 'CHAT_AI_MAX_UPLOAD_MB', 15) * 1024 * 1024
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f'Format {ext or "tanpa ekstensi"} belum didukung.')
    if uploaded.size > max_bytes:
        raise ValueError(f'Ukuran file maksimal {getattr(settings, "CHAT_AI_MAX_UPLOAD_MB", 15)} MB.')


def extract_text(file_path, original_name):
    ext = Path(original_name).suffix.lower()
    try:
        if ext in {'.txt', '.md'}:
            text = Path(file_path).read_text(encoding='utf-8', errors='replace')
        elif ext == '.csv':
            raw = Path(file_path).read_text(encoding='utf-8-sig', errors='replace')
            rows = list(csv.reader(io.StringIO(raw)))
            text = '\n'.join(' | '.join(cell for cell in row) for row in rows)
        elif ext == '.json':
            payload = json.loads(Path(file_path).read_text(encoding='utf-8', errors='replace'))
            text = json.dumps(payload, ensure_ascii=False, indent=2)
        elif ext == '.pdf':
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            text = '\n\n'.join((page.extract_text() or '') for page in reader.pages)
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            chunks = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    chunks.append(' | '.join(cell.text.strip() for cell in row.cells))
            text = '\n'.join(chunks)
        elif ext in {'.xlsx', '.xls'}:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            chunks = []
            for ws in wb.worksheets:
                chunks.append(f'### Sheet: {ws.title}')
                for row in ws.iter_rows(values_only=True):
                    values = ['' if v is None else str(v) for v in row]
                    if any(values):
                        chunks.append(' | '.join(values))
            text = '\n'.join(chunks)
        elif ext in {'.png', '.jpg', '.jpeg', '.webp'}:
            return '', ''
        else:
            return '', 'Format belum didukung untuk ekstraksi teks.'
        return text[:MAX_EXTRACTED_CHARS], ''
    except Exception as exc:
        return '', str(exc)
